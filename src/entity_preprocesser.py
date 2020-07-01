#!/usr/bin/env python
# coding: utf-8
import os
import re
import json
import spacy
import numpy as np
import pandas as pd
from docx import Document
from io import StringIO, BytesIO
from doc2text import para_to_text, sop_to_text

def read_sop(filepath):
    """
    Read SOP.docx files into plain text
    
    Params
    ----
    filepath: str    file path to SOP.docx
    
    Returns:
    obj    docx.Document object
    """
    with open(filepath, 'rb') as f:
        source_stream = BytesIO(f.read())
    f.close()
    doc = Document(source_stream)
    
    return doc


    

def event_parser(path_sop_folder="/Users/Public/Desktop/SOPs/"):
    """
    Parse Event Entity from SOPs
    
    Params
    ----
    path_sop_folder: str    folder path to where SOPs are stored
    path_sop_info: str    folder path to where SOP information are stored
    
    Returns:
    list    list of strings with spaCy Entity Ruler patterns
    """
    if not path_sop_folder.endswith("/"):
        path_sop_info = path_sop_info+"/"

    path = path_sop_folder
    SOPs = os.listdir(path)

    events = list()
    code_pattern = r"^[A-Z0-9][A-Z0-9]+[\- ]+.+"
    
    # Parsing event type from the SOP description part
    for sop in SOPs:
        filepath = path+sop
        doc = read_sop(filepath)

        paras = doc.paragraphs
        for p in paras:
            style = p.style.name
            text = para_to_text(p)

            if "Normal" in style:
                code = text.split(" event type")[0].split("for the ")[-1]
                r = re.findall(code_pattern, code)
                if r:
                    events.append(r[0])
                    
    # Parsing event type from the SOP filename (added Frank's code)
    pattern = r'^([A-Z]+)[ -]+([\w ]*)-([\w ,-]*).docx'
    events_code = list()
    
    for sop in SOPs:
        match = re.findall(pattern, sop, re.IGNORECASE)
        event_parts = [x.strip() for x in list(filter(None, match[0][1:]))]
        events_code.extend(event_parts)
        events_code.append(" - ".join(event_parts))

    # Remove redundant patterns
    events.extend(events_code)
    events = list(set(events))

    return events

def agency_parser():
    """
    Define a list of agencies
    """
    jurisdiction = ['AB', 'BI', 'BU', 'DE', 'DFPF', 
            'NW', 'PO', 'RI', 'RM', 'SC', 
            'SQ', 'SX', 'UN', 'VA', 'WP', 
            'WV', 'DF PF']
    
    return jurisdiction


def situation_parser(path_sop_folder="/Users/Public/Desktop/SOPs/"):
    """
    Parse Situation Entity from SOPs
    
    Params
    ----
    path_sop_folder: str    folder path to where SOPs are stored
    
    Returns:
    list    list of strings with spaCy Entity Ruler patterns
    """
    path = path_sop_folder
    SOPs = os.listdir(path)
    
    situations = list()

    for sop in SOPs:
        filepath = path+sop
        doc = read_sop(filepath)
        
        paras = doc.paragraphs
        for p in paras:
            style = p.style.name
            text = para_to_text(p)
            if "heading 2" in style.lower():
                situations.append(text)

    situations = list(set(situations))
    
    return situations
    
    
    
def role_parser(path_sop_folder="/Users/Public/Desktop/SOPs/"):
    """
    Parse Role Entity from SOPs
    
    Params
    ----
    path_sop_folder: str    folder path to where SOPs are stored
    
    Returns:
    list    list of strings with spaCy Entity Ruler patterns
    """
    path = path_sop_folder
    SOPs = os.listdir(path)

    roles = list()

    for sop in SOPs:
        filepath = path+sop
        doc = read_sop(filepath)
        
        paras = doc.paragraphs
        for p in paras:
            style = p.style.name
            text = para_to_text(p)
            if "heading 1" in style.lower():
                roles.append(text.strip())

    roles = list(set(roles))[1:]
    return roles

        
def organization_parser(path_sop_folder="/Users/Public/Desktop/SOPs/",
                        jargon_docx_path="/Users/jalee/Desktop/capstone-2020/utils/acronyms.docx"):
    """
    Parse Organization Entity from SOPs
    
    Params
    ----
    path_sop_folder: str    folder path to where SOPs are stored
    jargon_docx_path: str    file path to where acronyms.docx is stored
    
    Returns:
    list    list of strings with spaCy Entity Ruler patterns
    
    """
    doc = read_sop(jargon_docx_path)
    paras = doc.paragraphs
    
    # Set up empty dataframe
    df = pd.DataFrame(columns=["jargon","meaning"])
    jargon = []
    meaning = []
    
    # Extract acronym information from acronyms.docx
    for p in paras:
        parsed = p.text.strip().split("\t")
        if len(parsed) == 2:
            jargon.append(parsed[0].strip())
            meaning.append(parsed[1].strip())

    tables = doc.tables
    for t in tables:
        for r in t.rows:
            i = 0
            for c in r.cells:
                if i % 2 == 0:
                    i += 1
                    jargon.append(c.text.strip())
                else:
                    meaning.append(c.text.strip())
    
    # Fill in the empty dataframe
    df["jargon"] = jargon
    df["meaning"] = meaning                
    df["jargon"].replace("",np.nan, inplace=True)
    df["meaning"].replace("",np.nan, inplace=True)
    df = df.dropna().reset_index()[["jargon","meaning"]]

    terms = list(df["jargon"])
    definitions = list(df["meaning"])
    
    # Use Spacy to detect organization
    ## Load Spacy Pretrained model
    nlp = spacy.load("en_core_web_sm")

    ind = 0
    orgs = []
    
    # Cross check with the acronym definition to confirm that it is organization
    for d in definitions:
        doc = nlp(d)
        for ent in doc.ents:
            if ent.label_ == "ORG":
                orgs.append(terms[ind])
        ind += 1
    orgs = list(set(orgs))

    # Organization Names should be abbreviated capitals
    filtered = []
    for org in orgs:
        if org.upper() != org:
            pass
        else:
            filtered.append(org)

    org_cand = []

    glossary = zip(terms, definitions)
    for org, mean in glossary:
        if org in filtered:
            org_cand.append((org, mean))
    
    # Indicators of organization

    org_key = ["British","Columbia",
               "Service","Services",
               "Police","Institute",
               "Ltd.", "Association",
               "Ministry","National",
               "Unit","Incorporated",
               "Corporation"]

    organization = []

    for x,y in org_cand:
        if len(set(org_key + y.split())) < len(org_key + y.split()):
            organization.append(x)

    organization = list(set(organization))
    return organization



##### Incomplete ######
#######################

# def instruction_parser(path_sop_folder="/Users/Public/Desktop/SOPs/"):
#     """
#     Parse instruction (action, Condition, questions) relevant Entity from SOPs
   
#     Params
#     ----
#     path_sop_folder: str    folder path to where SOPs are stored
    
#     Returns:
#     list    list of strings with spaCy Entity Ruler patterns
#     """
#     path = path_sop_folder
#     SOPs = os.listdir(path)

#     raw_texts = []

#     for sop in SOPs:
#         filepath = path+sop
#         doc = read_sop(filepath)
        
#         paras = doc.paragraphs
#         for p in paras:
#             style = p.style.name
#             text = para_to_text(p)
#             if "style1" in style.lower():
#                 if text.strip() == "":
#                     pass
#                 else:
#                     raw_texts.append(text.strip())


#     questions = []
#     conditions = []
#     actions = []
#     others = []

#     for t in raw_texts:
#         if t.endswith("?"):
#             questions.append(t)
#         elif t.startswith("If"):
#             conditions.append(t)
#         elif nlp(t)[0].pos_ == "VERB":
#             actions.append(nlp(t)[0])
#         else:
#             others.append(t)

#     questions = list(set(questions))
#     conditions = list(set(conditions))
#     actions = list(set(actions))

#     filtered_conditions = []
#     for c in conditions:
#         filtered_conditions.append(c.replace("\\","").replace("/","or").replace("\t","").replace(":","").replace(";",""))
        
#     return
##### Incomplete ######
#######################



def entity_rule_writer(path_to_write="./entity_train/PATTERNS.JSONL", 
                        path_sop_folder="/Users/Public/Desktop/SOPs/",
                        jargon_docx_path="/Users/jalee/Desktop/capstone-2020/utils/acronyms.docx"):
    """
    Write PATTERNS.JSONL file for spaCy Entity Ruler model
    
    Params
    ----
    path_to_write: str    folder path to write the PATTERNS.JSONL file
    path_sop_folder: str    folder path to read the SOP docx files  
    """
    events = event_parser(path_sop_folder)
    situations = situation_parser(path_sop_folder)
    roles = role_parser(path_sop_folder)
    organizations = organization_parser(path_sop_folder, jargon_docx_path)
    jurisdiction = agency_parser()
#     questions = event_parser(path_sop_folder)
#     conditions = event_parser(path_sop_folder)
#     actions = event_parser(path_sop_folder)
    
    with open(path_to_write, mode="w", encoding="utf-8") as f:

        # JURISDICTION / AGENCY
        for j in jurisdiction:
            regex_pattern = "^({})[ -–]*".format(j)
            f.write('{"label":"JURI", "pattern":[{"TEXT":{"REGEX": "%s" }}]}\n' %regex_pattern)
        # EVENT CODE 
        for e in events:
            f.write('{"label":"EVENT", "pattern":"%s"}\n' %e)

        # SITUATION 
        for s in situations:
            f.write('{"label":"SITUATION", "pattern":"%s"}\n' %s)

        # ROLE 
        for r in roles:
            f.write('{"label":"ROLE", "pattern":"%s"}\n' %r)
        
        # ORGANIZATION
        for o in organizations:
            f.write('{"label":"ORG", "pattern":"%s"}\n' %o)

#         # QUESTION 
#         for q in questions:
#             f.write('{"label":"QUESTION", "pattern":"%s"}\n' %q)   

#         # CONDITION 
#         for c in filtered_conditions:
#             f.write('{"label":"CONDITION", "pattern":"%s"}\n' %c)  

#         # ACTION 
#         for a in actions:
#             f.write('{"label":"ACTION", "pattern":"%s"}\n' %a)    


    return
