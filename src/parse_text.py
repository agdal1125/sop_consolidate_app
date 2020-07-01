import os
import json
from docx import Document
from io import StringIO, BytesIO
import re
import pandas as pd

def main():
    """
    Convert SOP Docx files into csv files
    """
    def para_to_text(p):
        """
        A function to find every texts in the paragraph 
        (including the hypertext linked texts in blue)

        params
        ----
        p    docx.Document.Paragraph object

        returns
        ----
        str

        """
        rs = p._element.xpath(".//w:t")
        return u"".join([r.text for r in rs])




    def parse_docx_role(filepath):
        """
        Parse role-level SOP instructions from file

        params
        ----
        filepath: str    docx.Document.Paragraph object

        returns
        ----
        pd.dataframe 

        """
        with open(filepath, 'rb') as f:
            source_stream = BytesIO(f.read())
        f.close()
        doc = Document(source_stream)
        paras = doc.paragraphs
        res = dict()
        role = None
        texts, level = list(), list()

        # Iterate over paragraphs
        for p in paras:
            style = p.style.name
            text = para_to_text(p).strip()

            if len(text) == 0:
                continue
            if style == 'Heading 1':
                role = re.findall(r'([\w ]+)', text.lower())[0]
                res[role] = list()
            elif role is not None:
                res[role].append(text)
        
        return res




    def parse_docx_role_df(filepath):
        """
        Parse SOP Docx on role level

        params
        ----
        filepath: str    path to SOPs

        returns
        ----
        pd.DataFrame    dataframe that has information of sop on role levels (call taker, dispatcher)
        """
        with open(filepath, 'rb') as f:
            source_stream = BytesIO(f.read())
        f.close()
        doc = Document(source_stream)
        paras = doc.paragraphs
        res = dict()
        role = None
        texts, level = list(), list()
        for p in paras:
            style = p.style.name
            text = para_to_text(p)
            texts.append(text)
            level.append(style)
        return pd.DataFrame({'text': texts, 'level': level})




    def parse_one_event(curr_row):
        """
        Parse all available SOP documents for one event

        params
        ----
        curr_row:    pd.Series

        returns
        ----
        dict:    nested dictionary as a json 
        """
        cwd = os.getcwd()
        os.chdir(sop_dir)
        event_type = curr_row['type']
        jsons = dict()
        for juri, file in zip(curr_row['juri'], curr_row['filename']):
            jsons[juri] = parse_docx_role(file)
        os.chdir(cwd)
        return jsons




    def parse_docx_situation_df(doc_type, juri, filepath):
        """
        Parse SOP Docx on situation level

        params
        ----
        doc_type    pd.Series    event type of the SOP Document from 
        juri    pd.Series    jurisdiction of the SOP document from 
        filepath    path to SOPs

        returns
        ----
        pd.DataFrame:    dataframe that has information of SOP on situation levels (call taker, dispatcher)
        """
        with open(sop_dir + filepath, 'rb') as f:
            source_stream = BytesIO(f.read())
        f.close()
        doc = Document(source_stream)
        paras = doc.paragraphs
        res = pd.DataFrame()
        role, situation = None, ''
        roles, situations, texts = list(), list(), list()
        types, juris, files = list(), list(), list()

        for p in paras:
            style = p.style.name
            text = para_to_text(p)

            if len(text) == 0:
                continue
            if style == 'Normal':

                continue
            elif style == 'Heading 1':
                role = re.findall(r'([\w ]+)', text.lower())[0].strip()
                situation = ''
            elif style == 'Heading 2':
                situation = text.strip()
            elif role is not None:
                text = text.strip()
                roles.append(role)
                situations.append(situation)
                texts.append(text)

        res = pd.DataFrame({'role': roles, 'situation': situations, 'sop': texts}) \
                .groupby(['role', 'situation']) \
                .agg({'sop': lambda x: list(x)}) \
                .reset_index()
        res['juri'] = juri
        res['type'] = doc_type
        res['filename'] = filepath
        res = res[['type', 'juri', 'role', 'situation', 'sop', 'filename']]
        return res




    def get_all_situations(df):
        """
        Create a dataframe

        Params
        ----
        df    pd.DataFrame

        Returns
        ----
        pd.DataFrame:    dataframe that contains SOP information about event type, filename, jurisdiction by respective situations
        """
        # Create Empty Dataframe
        df_all_situation = pd.DataFrame()
        for i in range(len(df)):
            row = df.iloc[i, :]
            doc_type = row['type']
            for juri, filename in zip(row['juri'], row['filename']):
                situ_df = parse_docx_situation_df(doc_type, juri, filename)
                df_all_situation = df_all_situation.append(situ_df)
        return df_all_situation.reset_index(drop = True)




    # Filtering SOPs for Analysis
    ## Get list of raw SOPs
    curr_dir = os.getcwd()
    sop_dir="/Users/Public/Desktop/SOPs/"
    csv_dir="data/interim/"
    sop_list = os.listdir(sop_dir)

    ## Define desired jurisdictions and unwanted files
    juri_list = {'AB', 'BI', 'BU', 'DE', 'DFPF', 
                'NW', 'PO', 'RI', 'RM', 'SC', 
                'SQ', 'SX', 'UN', 'VA', 'WP', 
                'WV', 'DF PF'}
    
    ## SOPs to exclude from analysis (Fire Department, Non-police SOPs)
    exclude_docs = ['- report' # gray; non-police
                    , '- drugs -' # gray; non-police
                    , 'bomb threats' # red; fire-related
                    , 'water problem' # red; fire-related
                    , 'wildfire'
                    , 'stove fire'
                    , 'electrical problem'
                   ]

    juris, types, fullname = list(), list(), list()
    excl_juris, excl_types, excl_docs, excl_fullname = list(), list(), list(), list()
    unmatched_sops = list()

    ## Regex for parsing information from SOP file names
    sopreg = r'^([A-Z]+)[ -]+([\w ]*)-([\w ,]*).docx'

    ## Filter SOP with `exclude_docs`
    for sop in sop_list:
        isExcluded = re.search(r'|'.join(exclude_docs), sop, re.IGNORECASE)
        if isExcluded:
            unmatched_sops.append(sop)
            continue

        match = re.findall(sopreg, sop, re.IGNORECASE)
        if not match:
            unmatched_sops.append(sop)
            continue

        juri, acrn, doc = match[0]
        juri = re.sub(' ', '', juri.upper())
        doc = doc.upper().strip()
        acrn = acrn.upper().strip()

        ### Find inconsistency in SOP filenames
        if juri not in juri_list:
            excl_juris.append(juri)
            excl_types.append(acrn)
            excl_docs.append(doc)
            excl_fullname.append(sop)
            continue

        type_adjust = doc if len(acrn) == 0 else acrn
        juris.append(juri)
        types.append(type_adjust)
        fullname.append(sop)
    
    sop_docs = pd.DataFrame({'juri': juris, 'type': types, 'filename': fullname})
    sop_unmatched = pd.DataFrame({'filename': unmatched_sops})
    sop_excludedjuri = pd.DataFrame({'juri': excl_juris, 'type': excl_docs, 'filename': excl_fullname})
    
    type_to_juri = sop_docs.groupby('type') \
                    .agg({'juri': lambda x: list(x), 
                          'filename': lambda x: list(x)}) \
                    .reset_index()
    type_to_juri.loc[:, 'juri_count'] = type_to_juri.loc[:, 'juri'].apply(len)
    type_to_juri = type_to_juri.sort_values(by = 'juri_count', ascending = False, ignore_index = True)
    
    type_to_juri_ok = type_to_juri[type_to_juri['juri_count'] >= 13]

    ## Information of Filtered SOPs into csv files

    ## Information of Valid SOPs for analysis
    type_to_juri_ok.to_csv('data/interim/sop_types_valid.csv', index = False)

    ## Information of Excluded SOPs
    type_to_juri_excluded = type_to_juri[type_to_juri['juri_count'] < 13]
    type_to_juri_excluded.to_csv('data/interim/sop_types_excluded.csv', index = False)

    ## Information of SOP excluded due to Jurisdiction
    sop_excludedjuri.to_csv('data/interim/sop_excluded_due_to_juris.csv', index = False)

    ## Information of SOP excluded because of event type
    sop_unmatched.to_csv('data/interim/sop_excluded_due_to_type.csv', index = False)




    # convert SOP to Json file
    ## save json files to sop_jsons folder
    num_rows = type_to_juri_ok.shape

    for i in range(num_rows[0]):
        row = type_to_juri_ok.iloc[i, :]
        event = parse_one_event(row)
        doctype = row['type']
        outfile_name = f'data/sop_jsons/{doctype}.txt'
        with open(outfile_name, 'w') as outfile:
            json.dump(event, outfile)
        outfile.close()




    # Convert filtered SOPs into csv files

    ## All situation
    df_all_situation = get_all_situations(type_to_juri_ok)
    df_all_situation.to_csv('data/interim/all_situation.csv', index = False)
    df_call = df_all_situation[ 
        (df_all_situation['role'] == 'call taker') 
        & (df_all_situation['situation'].str.len() > 0)
    ].reset_index(drop = True)

    ## Call taker parts of SOP
    df_call.to_csv('data/interim/calltaker_situation.csv', index = False)

    ## Dispatcher parts of SOP
    df_dispatcher = df_all_situation[ 
        (df_all_situation['role'] == 'dispatcher') 
        & (df_all_situation['situation'].str.len() > 0)
    ].reset_index(drop = True)
    df_call.to_csv('data/interim/dispatcher_situation.csv', index = False)

if __name__ == "__main__":
    main()